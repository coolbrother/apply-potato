"""
Tests for docx_utils — resume reading and cover letter writing.

Usage:
    pytest tests/test_docx_utils.py -v
"""

import pytest
from pathlib import Path
from docx import Document

from src.docx_utils import read_resume_text, write_simple_cover_letter


@pytest.fixture
def tmp_resume(tmp_path) -> Path:
    """Create a minimal .docx resume for testing."""
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane@example.com | github.com/jane")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Software Engineer — Acme Corp, 2022–2024")
    doc.add_paragraph("Skills")
    # Add a table (common in two-column resumes)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].paragraphs[0].text = "Python"
    table.rows[0].cells[1].paragraphs[0].text = "TypeScript"
    path = tmp_path / "resume.docx"
    doc.save(str(path))
    return path


class TestReadResumeText:

    def test_reads_paragraphs(self, tmp_resume):
        text = read_resume_text(tmp_resume)
        assert "Jane Doe" in text
        assert "Software Engineer" in text

    def test_reads_table_cells(self, tmp_resume):
        text = read_resume_text(tmp_resume)
        assert "Python" in text
        assert "TypeScript" in text

    def test_no_duplicates(self, tmp_resume):
        text = read_resume_text(tmp_resume)
        # Each line should appear only once
        assert text.count("Jane Doe") == 1

    def test_returns_string(self, tmp_resume):
        assert isinstance(read_resume_text(tmp_resume), str)

    def test_short_resume_logs_warning(self, tmp_path, caplog):
        doc = Document()
        doc.add_paragraph("Short")
        path = tmp_path / "short.docx"
        doc.save(str(path))
        import logging
        with caplog.at_level(logging.WARNING, logger="src.docx_utils"):
            read_resume_text(path)
        assert "very short" in caplog.text


class TestWriteSimpleCoverLetter:

    def test_creates_file(self, tmp_path):
        out = tmp_path / "cl.docx"
        write_simple_cover_letter(out, "Dear Hiring Manager,\n\nI am excited.", "Jane Doe", "Acme", "SWE")
        assert out.exists()

    def test_output_path_returned(self, tmp_path):
        out = tmp_path / "cl.docx"
        result = write_simple_cover_letter(out, "Body text.", "Jane Doe", "Acme", "SWE")
        assert result == out

    def test_contains_user_name(self, tmp_path):
        out = tmp_path / "cl.docx"
        write_simple_cover_letter(out, "Body text.", "Jane Doe", "Acme", "SWE")
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in full_text

    def test_contains_company(self, tmp_path):
        out = tmp_path / "cl.docx"
        write_simple_cover_letter(out, "Body text.", "Jane Doe", "Acme Corp", "SWE")
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Acme Corp" in full_text

    def test_contains_position(self, tmp_path):
        out = tmp_path / "cl.docx"
        write_simple_cover_letter(out, "Body text.", "Jane Doe", "Acme", "Senior Engineer")
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Senior Engineer" in full_text

    def test_body_paragraphs_included(self, tmp_path):
        out = tmp_path / "cl.docx"
        body = "First paragraph here.\n\nSecond paragraph here."
        write_simple_cover_letter(out, body, "Jane Doe", "Acme", "SWE")
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "First paragraph here." in full_text
        assert "Second paragraph here." in full_text
