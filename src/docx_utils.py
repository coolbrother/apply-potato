"""
Utilities for reading and writing .docx files.
"""

import logging
from datetime import date
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def read_resume_text(path: Path) -> str:
    """
    Extract plain text from a .docx resume file.

    Reads paragraph text and table cell text (for two-column resume layouts).
    Logs a warning if the extracted text is suspiciously short.
    """
    from docx import Document  # lazy import — only needed when resume reading is required

    doc = Document(str(path))
    lines = []
    seen = set()

    for p in doc.paragraphs:
        text = p.text.strip()
        if text and text not in seen:
            lines.append(text)
            seen.add(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text and text not in seen:
                        lines.append(text)
                        seen.add(text)

    result = "\n".join(lines)
    if len(result) < 200:
        logger.warning(f"Resume text extracted from {path} is very short ({len(result)} chars) — verify the file is correct")
    return result


def write_simple_cover_letter(
    output_path: Path,
    cover_letter_text: str,
    user_name: str,
    company: str,
    position: str,
) -> Path:
    """
    Write a plain cover letter docx for non-claude providers.

    Uses simple paragraph formatting only — no XML manipulation.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Default style: Calibri 11pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    def add_para(text: str = "", bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold

    # Header: name
    add_para(user_name, bold=True)
    # Date
    add_para(date.today().strftime("%B %d, %Y"))
    add_para()
    # Company and role
    add_para(company)
    add_para(f"Re: {position}")
    add_para()

    # Body paragraphs (split on double newline)
    for paragraph in cover_letter_text.strip().split("\n\n"):
        text = paragraph.strip()
        if text:
            add_para(text)
            add_para()

    doc.save(str(output_path))
    logger.debug(f"Cover letter written to {output_path}")
    return output_path
